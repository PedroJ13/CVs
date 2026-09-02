from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"C:\Work\CVs\Output\Pedro_Gutierrez_INDI_SQL_Developer_Senior_CV.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9E2EC", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, name="Calibri", size=10.2, color="1F2937", bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    set_run_font(r, size=10.5, color="1F4D78", bold=True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(2.2)
    p.paragraph_format.line_spacing = 1.03
    r = p.add_run(text)
    set_run_font(r, size=9.2)
    return p


def add_role(doc, title, company, dates, location, bullets, tech=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{company} - {title}")
    set_run_font(r, size=10.2, color="111827", bold=True)
    r = p.add_run(f" | {dates} | {location}")
    set_run_font(r, size=9.2, color="4B5563")
    for bullet in bullets:
        add_bullet(doc, bullet)
    if tech:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("Technologies: ")
        set_run_font(r, size=8.8, color="374151", bold=True)
        r = p.add_run(tech)
        set_run_font(r, size=8.8, color="374151")


def add_skill_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    set_table_borders(table)
    set_cell_margins(table)
    for i, (label, value) in enumerate(rows):
        cells = table.rows[i].cells
        set_cell_width(cells[0], 2050)
        set_cell_width(cells[1], 7310)
        set_cell_shading(cells[0], "F2F4F7")
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, size=8.8, color="1F4D78", bold=True)
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, size=8.8, color="1F2937")


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.65)
section.right_margin = Inches(0.65)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(9.5)

for style_name in ("List Bullet", "List Paragraph"):
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(9.2)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(1)
r = title.add_run("PEDRO JAVIER GUTIERREZ ARMAS")
set_run_font(r, size=16, color="0B2545", bold=True)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(1)
r = subtitle.add_run("Senior SQL Server Developer | T-SQL, SSIS, Stored Procedures, Automation & Performance Tuning")
set_run_font(r, size=10.4, color="1F4D78", bold=True)

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.paragraph_format.space_after = Pt(5)
r = contact.add_run("San Jose, Costa Rica | +506 6351-5860 | pj13eros@hotmail.com | linkedin.com/in/pedrogutierrez13")
set_run_font(r, size=8.8, color="374151")

add_heading(doc, "Professional Summary")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
p.paragraph_format.line_spacing = 1.05
summary = (
    "SQL Server Developer, DBA, and Data Engineer with 15+ years of experience designing, developing, "
    "optimizing, troubleshooting, and maintaining database and data integration solutions. Strong hands-on "
    "background with SQL Server, T-SQL, stored procedures, complex queries, database objects, SSIS, SSRS, SSAS, "
    "Power BI, Python, PowerShell, Azure-based environments, PostgreSQL, MySQL, Snowflake SQL, and production "
    "data workflows. Experienced resolving production database issues, improving query performance, preparing "
    "schema and data scripts, automating data processing, integrating disparate data sources, and collaborating "
    "with engineering and business teams in English-speaking environments."
)
r = p.add_run(summary)
set_run_font(r, size=9.4)

add_heading(doc, "Technical Skills")
add_skill_table(
    doc,
    [
        ("SQL Server Development", "T-SQL, stored procedures, functions, views, complex queries, database objects, schemas, table structures, indexes."),
        ("SSIS & Integration", "SSIS packages, ETL/ELT processes, data migration, multi-source consolidation, Salesforce, MySQL, PostgreSQL, Snowflake SQL."),
        ("Production Support", "Production troubleshooting, query diagnostics, error handling, validation checks, data integrity, migration continuity."),
        ("Performance", "Query optimization, indexing strategies, execution-time reduction, response-time improvement, database workflow tuning."),
        ("Automation & Scripting", "Python, PowerShell, Pandas, NumPy, Matplotlib, automated validation, reporting workflow automation."),
        ("Microsoft Data Stack", "SQL Server, SSIS, SSAS, SSRS, Power BI, Excel, Azure-based environments."),
    ],
)

add_heading(doc, "Professional Experience")
add_role(
    doc,
    "Snowflake Developer / Data Engineer",
    "ServiceTitan",
    "Sep 2024 - Present",
    "Remote / Costa Rica",
    [
        "Migrate reporting and business logic from C# processes into Snowflake SQL, improving maintainability and traceability of analytics workflows.",
        "Create and maintain dbt models within silver and gold data layers to support reusable reporting datasets and downstream analytics.",
        "Optimize Snowflake SQL and dbt workloads, reducing data warehouse processing time by approximately 20% in an initial optimization phase.",
        "Support MetricFlow semantic layer work with Snowflake Cortex, AI-assisted development tools, Cursor, and PR review automation.",
    ],
    "Snowflake, Snowflake SQL, dbt, MetricFlow, Snowflake Cortex, SQL, C# logic migration, Cursor, AI-assisted PR review bot",
)

add_role(
    doc,
    "Data Engineer",
    "SMASH Costa Rica",
    "May 2021 - Sep 2024",
    "San Jose, Costa Rica",
    [
        "Designed and maintained ETL workflows using SQL Server, SSIS, Python, PowerShell, Snowflake, Power BI, and Excel, reducing manual processing time by up to 40%.",
        "Built automated data refresh, validation, and exploratory analysis workflows to detect anomalies and improve data validation accuracy by 30%.",
        "Prepared, cleaned, transformed, and modeled datasets for business reporting, Power BI dashboards, and operational decision-making.",
        "Collaborated with stakeholders and technical teams to translate reporting needs into repeatable data routines and analytics-ready outputs.",
    ],
    "SQL Server, SSIS, T-SQL, Python, PowerShell, Snowflake, Power BI, Excel, Pandas, NumPy, Matplotlib",
)

add_role(
    doc,
    "SQL Developer",
    "Intertec International",
    "Feb 2019 - Apr 2021",
    "San Jose, Costa Rica",
    [
        "Developed and optimized SQL Server stored procedures, complex T-SQL queries, and database workflows supporting business logic and analytics needs.",
        "Reduced query execution times by up to 50% through SQL optimization, indexing strategies, and performance tuning.",
        "Analyzed and resolved data workflow issues using validation checks, error handling, and structured troubleshooting across key processes.",
        "Integrated disparate data sources, including Salesforce and MySQL, into analytics-ready datasets using optimized ETL processes.",
    ],
    "SQL Server, T-SQL, SSIS, stored procedures, Salesforce, MySQL, query optimization, indexing",
)

add_role(
    doc,
    "DBA / SQL Developer",
    "EL Tiempo",
    "Sep 2020 - Nov 2020",
    "Colombia",
    [
        "Supported migration projects for 10 databases, coordinating validation, data integrity checks, risk identification, and continuity activities.",
        "Prepared and validated database migration routines to preserve reporting availability and operational continuity during deployment activities.",
        "Worked with cross-functional teams to diagnose migration risks, validate data, and maintain database integrity.",
    ],
    "SQL Server, SSIS, Azure, SSRS",
)

add_role(
    doc,
    "DBA / SQL & BI Consultant",
    "Gold Data Networks",
    "Jan 2016 - Feb 2020",
    "Panama City, Panama",
    [
        "Built relational databases, table structures, custom database objects, and stored procedures from the ground up for application and reporting workloads.",
        "Designed end-to-end SQL Server, SSIS, SSRS, PostgreSQL, Power BI, and Excel solutions integrated with existing infrastructure and BI needs.",
        "Improved backend and reporting performance through database object design, stored procedure development, and SQL Server best practices.",
    ],
    "SQL Server, T-SQL, SSIS, SSRS, PostgreSQL, Power BI, Excel",
)

add_role(
    doc,
    "Data Warehouse DBA",
    "BAC Credomatic",
    "Nov 2017 - Jan 2019",
    "San Jose, Costa Rica",
    [
        "Developed and maintained ETL pipelines to extract, transform, and load data from multiple sources into the data warehouse.",
        "Defined and optimized table, index, and view structures for high-performance analytical workloads.",
        "Supported SSAS, Power BI, SSRS, and Azure-based reporting environments for large-scale analytics and decision support.",
    ],
    "SQL Server, SSIS, SSAS, SSRS, Power BI, Azure, data warehouse",
)

add_heading(doc, "Additional SQL Server Experience")
for line in [
    "Database Manager / SQL & BI roles across Xetux Solutions, ACH Cloud Services, EducaTablet, VIGEOSOFT, and Optica Caroni C.A.",
    "Administered and developed SQL Server databases, reporting workflows, SSIS/SSRS assets, database policies, data dictionaries, and production process improvements.",
    "Modeled, designed, configured, and programmed OLTP databases and trained development teams on database programming good practices.",
]:
    add_bullet(doc, line)

add_heading(doc, "Education & Certifications")
for line in [
    "Systems Analyst, Informatics - IUT Dr. Federico Rivero Palacio.",
    "English Certificate - Universidad Central de Venezuela / Microsoft.",
    "SQL Admin Part 1.",
    "Analyzing and Visualizing Data with Power BI.",
    "Complete Data Science Training with Python for Data Analysis.",
    "R Programming A-Z: R for Data Science.",
    "Dataiku Core Designer.",
]:
    add_bullet(doc, line)

add_heading(doc, "Languages")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
r = p.add_run("Spanish: Native | English: Full professional proficiency")
set_run_font(r, size=9.3)

doc.core_properties.author = "Pedro Javier Gutierrez Armas"
doc.core_properties.title = "Pedro Gutierrez - INDI SQL Developer Senior CV"
doc.save(OUTPUT)
print(OUTPUT)
