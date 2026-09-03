from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"C:\Work\CVs\Output\Pedro_Gutierrez_CCS_Sr_Data_Engineer_MS_Stack_CV.docx"


def set_run_font(run, name="Calibri", size=9.4, color="1F2937", bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9E2EC", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn("w:" + name))
        if node is None:
            node = OxmlElement("w:" + name)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text.upper())
    set_run_font(r, size=10.5, color="1F4D78", bold=True)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(2.2)
    p.paragraph_format.line_spacing = 1.03
    r = p.add_run(text)
    set_run_font(r, size=9.15)


def add_role(doc, company, title, dates, location, bullets, technologies, page_break_before=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = page_break_before
    r = p.add_run(f"{company} - {title}")
    set_run_font(r, size=10.1, color="111827", bold=True)
    r = p.add_run(f" | {dates} | {location}")
    set_run_font(r, size=9.1, color="4B5563")
    for item in bullets:
        add_bullet(doc, item)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Technologies: ")
    set_run_font(r, size=8.8, color="374151", bold=True)
    r = p.add_run(technologies)
    set_run_font(r, size=8.8, color="374151")


def add_skills_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    set_table_borders(table)
    set_cell_margins(table)
    for i, (label, value) in enumerate(rows):
        left, right = table.rows[i].cells
        set_cell_width(left, 2100)
        set_cell_width(right, 7260)
        shade(left, "E8EEF5")
        left.paragraphs[0].paragraph_format.space_after = Pt(0)
        right.paragraphs[0].paragraph_format.space_after = Pt(0)
        r = left.paragraphs[0].add_run(label)
        set_run_font(r, size=8.7, color="1F4D78", bold=True)
        r = right.paragraphs[0].add_run(value)
        set_run_font(r, size=8.7, color="1F2937")


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.65)
section.right_margin = Inches(0.65)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(9.5)

for style_name in ("List Bullet", "List Paragraph"):
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(9.15)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(1)
r = title.add_run("PEDRO JAVIER GUTIERREZ ARMAS")
set_run_font(r, size=16, color="0B2545", bold=True)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(1)
r = subtitle.add_run("Sr Data Engineer | Microsoft Data Stack, SQL Server, SSIS, Power BI, Snowflake & Reporting")
set_run_font(r, size=10.25, color="1F4D78", bold=True)

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.paragraph_format.space_after = Pt(5)
r = contact.add_run("San Jose, Costa Rica | +506 6351-5860 | pj13eros@hotmail.com | linkedin.com/in/pedrogutierrez13")
set_run_font(r, size=8.8, color="374151")

add_heading(doc, "Professional Summary")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
p.paragraph_format.line_spacing = 1.05
summary = (
    "Senior Data Engineer, Snowflake Developer, and SQL Server specialist with 15+ years of experience "
    "designing, building, optimizing, and maintaining data solutions across ETL/ELT pipelines, data warehousing, "
    "BI/reporting, SQL development, data quality, and production support. Strong hands-on background with SQL Server, "
    "T-SQL, SSIS, SSAS, SSRS, Power BI, Excel, Python, PowerShell, Azure-based environments, Snowflake SQL, dbt, "
    "PostgreSQL, MySQL, and Salesforce integrations. Experienced building reporting-ready datasets, integrating "
    "multiple data sources, automating data processing, optimizing warehouse workloads, and collaborating with Data "
    "Engineering and BI stakeholders in remote English-speaking environments."
)
r = p.add_run(summary)
set_run_font(r, size=9.3)

add_heading(doc, "Technical Skills")
add_skills_table(
    doc,
    [
        ("Data Engineering", "ETL/ELT pipelines, data transformation, data integration, data warehousing, data modeling, silver/gold data layers."),
        ("Microsoft Stack", "SQL Server, T-SQL, SSIS, SSAS, SSRS, Power BI, Excel, Azure-based environments, reporting datasets."),
        ("Cloud Data Platforms", "Snowflake SQL, dbt models, Snowflake Cortex, MetricFlow semantic layer support, PostgreSQL, MySQL."),
        ("Automation", "Python, PowerShell, Pandas, NumPy, Matplotlib, automated validation, refresh workflows, anomaly detection."),
        ("Performance & Quality", "Query optimization, indexing strategies, warehouse processing optimization, validation checks, error handling."),
        ("Delivery", "Stakeholder collaboration, requirements translation, BI enablement, process documentation, production data support."),
    ],
)

add_heading(doc, "Professional Experience")
add_role(
    doc,
    "ServiceTitan",
    "Snowflake Developer / Data Engineer",
    "Sep 2024 - Present",
    "Remote / Costa Rica",
    [
        "Migrate reporting and business logic from C# processes into Snowflake SQL to support scalable reporting and analytics workflows.",
        "Create and maintain dbt models in silver and gold data layers, improving model organization, reusability, and downstream reporting reliability.",
        "Optimize Snowflake SQL and dbt workloads, reducing data warehouse processing time by approximately 20% in an initial optimization phase.",
        "Support MetricFlow semantic layer work with Snowflake Cortex, improving metric organization and consistency for analytics consumers.",
    ],
    "Snowflake, Snowflake SQL, dbt, MetricFlow, Snowflake Cortex, SQL, C# logic migration, Cursor, AI-assisted PR review bot",
)
add_role(
    doc,
    "SMASH Costa Rica",
    "Data Engineer",
    "May 2021 - Sep 2024",
    "San Jose, Costa Rica",
    [
        "Designed and implemented customized ETL applications tailored to client requirements, reducing manual processing time by up to 40%.",
        "Built and maintained data workflows with SQL Server, SSIS, Python, PowerShell, Snowflake, Power BI, and Excel for reporting and analytics.",
        "Developed automated exploratory data analysis and validation tools with Pandas, NumPy, and Matplotlib to detect anomalies and improve validation accuracy by 30%.",
        "Prepared, cleaned, transformed, and modeled datasets for Power BI dashboards, operational reporting, and business decision support.",
    ],
    "SQL Server, SSIS, Power BI, Python, Pandas, NumPy, Matplotlib, PowerShell, Snowflake, Excel",
)
add_role(
    doc,
    "Intertec International",
    "SQL Developer",
    "Feb 2019 - Apr 2021",
    "San Jose, Costa Rica",
    [
        "Developed and optimized stored procedures, SQL queries, and database workflows supporting complex business logic and analytics needs.",
        "Consolidated multiple disparate data sources into analytics-ready datasets using optimized ETL processes.",
        "Reduced query execution times by up to 50% through SQL optimization, indexing strategies, and performance tuning.",
        "Improved data accuracy by over 25% through validation checks and error-handling mechanisms across key workflows.",
    ],
    "SQL Server, T-SQL, SSIS, Salesforce, MySQL, stored procedures, query optimization, indexing",
)
add_role(
    doc,
    "EL Tiempo",
    "DBA / SQL Developer",
    "Sep 2020 - Nov 2020",
    "Colombia",
    [
        "Led planning and execution of migration activities for 10 databases, supporting data validation, integrity, and continuity of operations.",
        "Worked with cross-functional teams to identify migration risks, validate data, and maintain reporting availability during deployment activities.",
        "Supported SQL Server, SSIS, Azure, and SSRS assets in a multi-platform migration and reporting environment.",
    ],
    "SQL Server, SSIS, Azure, SSRS",
    page_break_before=True,
)
add_role(
    doc,
    "Gold Data Networks",
    "DBA / SQL & BI Consultant",
    "Jan 2016 - Feb 2020",
    "Panama City, Panama",
    [
        "Built relational databases, table structures, custom database objects, and stored procedures from the ground up for application and reporting workloads.",
        "Delivered SQL Server, SSIS, SSRS, PostgreSQL, Power BI, and Excel solutions integrated with existing infrastructure and BI needs.",
        "Designed end-to-end data solutions to improve reporting, database performance, and scalable backend processing.",
    ],
    "SQL Server, T-SQL, SSIS, SSRS, PostgreSQL, Power BI, Excel",
)
add_role(
    doc,
    "BAC Credomatic",
    "Data Warehouse DBA",
    "Nov 2017 - Jan 2019",
    "San Jose, Costa Rica",
    [
        "Developed and maintained ETL pipelines to extract, transform, and load data from multiple sources into the data warehouse.",
        "Defined and optimized table, index, and view structures for high-performance analytical workloads.",
        "Supported SSAS, Power BI, SSRS, and Azure-based reporting environments for large-scale analytics and decision support.",
    ],
    "SQL Server, SSIS, SSAS, Power BI, SSRS, Azure, data warehouse",
)

add_heading(doc, "Additional Data & Microsoft Stack Experience")
for item in [
    "Built the first stage of a main data warehouse and generated reports and presentations for senior management at Bosal.",
    "Created a data lake to consolidate sales data and support centralized reporting and analysis at Xetux Solutions.",
    "Administered and developed SQL Server databases, SSIS/SSRS assets, dashboards, billing reports, data dictionaries, and OLTP database processes across ACH Cloud Services, EducaTablet, VIGEOSOFT, and Optica Caroni C.A.",
]:
    add_bullet(doc, item)

add_heading(doc, "Education & Certifications")
for item in [
    "Systems Analyst, Informatics - IUT Dr. Federico Rivero Palacio.",
    "English Certificate - Universidad Central de Venezuela / Microsoft.",
    "SQL Admin Part 1.",
    "Analyzing and Visualizing Data with Power BI.",
    "Complete Data Science Training with Python for Data Analysis.",
    "R Programming A-Z: R for Data Science.",
    "Dataiku Core Designer.",
]:
    add_bullet(doc, item)

add_heading(doc, "Languages")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
r = p.add_run("Spanish: Native | English: Full professional proficiency")
set_run_font(r, size=9.3)

doc.core_properties.author = "Pedro Javier Gutierrez Armas"
doc.core_properties.title = "Pedro Gutierrez - CCS Sr Data Engineer Microsoft Data Stack CV"
doc.save(OUTPUT)
print(OUTPUT)
